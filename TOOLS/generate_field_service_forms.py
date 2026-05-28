import csv
import json
import shutil
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(r"C:\HVAC_PRO_MSE\research_downloads\field_service_forms")
ZIP_PATH = Path(r"C:\HVAC_PRO_MSE\research_downloads\ServoERP_Field_Service_Form_Templates.zip")
DATE_ACCESSED = datetime.now().strftime("%Y-%m-%d")
VERSION = "2.0.0-premium"


CATEGORIES = [
    ("01_HVAC", "HVAC", [
        "site survey form", "equipment inventory form", "inspection form", "preventive maintenance checklist",
        "breakdown/service call form", "service report", "commissioning checklist", "installation checklist",
        "startup/shutdown checklist", "filter cleaning checklist", "coil cleaning checklist", "refrigerant top-up record",
        "temperature/pressure readings sheet", "AHU checklist", "FCU checklist", "chiller checklist",
        "cooling tower checklist", "duct inspection form", "water treatment log", "AMC visit report",
        "spare parts usage form", "quote/estimate form", "customer sign-off form",
    ]),
    ("02_Plumbing", "Plumbing", [
        "site inspection form", "leak report form", "drainage inspection form", "pump check form",
        "valve/fixture checklist", "water tank cleaning checklist", "pressure test form", "pipe repair form",
        "blockage clearance report", "bathroom/kitchen fit-out checklist", "preventive maintenance checklist",
        "AMC visit report", "parts consumed form", "quotation form", "handover/sign-off form",
    ]),
    ("03_Electrical", "Electrical", [
        "electrical inspection form", "panel board preventive maintenance form", "load/voltage reading sheet",
        "breaker test checklist", "earthing test form", "insulation/resistance test form", "lighting inspection form",
        "emergency lighting checklist", "switchgear inspection form", "UPS/inverter checklist",
        "generator electrical checklist", "troubleshooting form", "safety isolation / lockout-tagout form",
        "installation commissioning form", "service report", "parts usage form", "customer approval/sign-off form",
    ]),
    ("04_Refrigeration", "Refrigeration", [
        "refrigeration inspection form", "cold room checklist", "compressor checklist",
        "evaporator/condenser inspection form", "refrigerant leak check form", "pressure/temperature log sheet",
        "defrost cycle checklist", "compressor start-up checklist", "maintenance checklist", "breakdown report",
        "AMC service report", "gas charging/recovery record", "spare parts issue form", "quotation form",
        "customer sign-off form",
    ]),
    ("05_Fire_Safety", "Fire Safety", [
        "fire system inspection form", "fire NOC inspection checklist", "extinguisher inspection form",
        "hydrant/sprinkler inspection form", "fire alarm panel checklist", "smoke detector checklist",
        "emergency lighting/exit sign checklist", "pump room checklist", "fire drill report",
        "evacuation drill record", "annual compliance audit form", "fire safety maintenance log", "defect report",
        "corrective action report", "completion certificate/sign-off form",
    ]),
    ("06_Commercial_Service_Shared", "Commercial Service Shared", [
        "lead/enquiry form", "site visit form", "estimate/quotation form", "work order form", "job card",
        "preventive maintenance schedule", "breakdown service report", "service completion report",
        "asset history form", "customer feedback form", "complaint form", "invoice form", "payment receipt form",
        "warranty/AMC form",
    ]),
    ("07_Dispatch_Office", "Dispatch Office", [
        "dispatch assignment form", "technician attendance form", "leave/request form",
    ]),
    ("08_Inventory", "Inventory", [
        "spare parts requisition form", "stock issue/return form", "purchase request form", "goods received note",
    ]),
    ("09_Finance_Accounting", "Finance Accounting", [
        "invoice approval form", "credit note form", "collections follow-up form", "job costing sheet",
    ]),
    ("10_Compliance_Safety", "Compliance Safety", [
        "site risk assessment form", "PPE inspection checklist", "toolbox talk record", "incident/near-miss report",
        "permit to work checklist", "working at height checklist", "hot work permit checklist",
        "confined space entry checklist",
    ]),
]


SEARCH_REFERENCES = [
    ("HVAC preventive maintenance checklist PDF public", "https://mechbasic.com/wp-content/uploads/2025/04/HVAC_Preventive_Maintenance_Checklist.pdf", "Public web PDF found; not copied to avoid unclear reuse/licensing and wording reuse."),
    ("HVAC preventive maintenance checklist PDF public", "https://www.camcode.com/wp-content/uploads/2021/09/hvac-preventive-maintenance-checklist.pdf", "Public web PDF found; not copied to avoid branded/proprietary template wording."),
    ("fire extinguisher inspection checklist PDF public", "https://ehs.berkeley.edu/publications/fire-extinguisher-monthly-inspection-guide", "Public university safety guidance found; used only as source-awareness, not copied."),
    ("fire extinguisher inspection checklist PDF public", "https://compliancetrainingpartners.com/wp-content/uploads/2021/03/Fire-Extinguisher-Inspection-Checklist.pdf", "Public web PDF found; not copied to avoid template ownership ambiguity."),
    ("electrical inspection checklist PDF India", "", "Search query retained for manual review; no clean unbranded legal direct template selected."),
    ("panel maintenance checklist Excel", "", "Search query retained for manual review; no clean unbranded legal direct template selected."),
    ("plumbing pressure test report template", "", "Search query retained for manual review; no clean unbranded legal direct template selected."),
    ("cold room maintenance checklist", "", "Search query retained for manual review; no clean unbranded legal direct template selected."),
    ("goods received note format Excel India", "", "Search query retained for manual review; no clean unbranded legal direct template selected."),
]


TRADE_PROFILES = {
    "HVAC": {
        "assets": ["Split AC", "VRV/VRF", "AHU", "FCU", "Chiller", "Cooling Tower", "Ducting", "Pump Set"],
        "readings": [
            ("Return air temperature", "C", "18-30"), ("Supply air temperature", "C", "8-18"),
            ("Delta T", "C", "8-14"), ("Suction pressure", "psi/bar", "OEM range"),
            ("Discharge pressure", "psi/bar", "OEM range"), ("Compressor current", "A", "Nameplate +/-10%"),
            ("Refrigerant type", "text", "R32/R410A/R134a/etc."), ("Drain flow", "Pass/Fail", "Clear"),
            ("Filter differential", "Pa", "Site standard"), ("BMS/thermostat status", "Pass/Fail", "Operational"),
        ],
        "tasks": [
            "Inspect air filters and clean/replace as required", "Inspect evaporator/condenser coils",
            "Check fan motor, blower, belt, pulley and vibration", "Inspect drain tray and condensate line",
            "Verify thermostat/BMS command and feedback", "Record refrigerant/pressure/current readings",
            "Check electrical terminals and insulation condition", "Confirm cooling/heating performance",
        ],
        "risks": ["Refrigerant handling", "Electrical isolation", "Working at height", "Condensate leakage"],
    },
    "Plumbing": {
        "assets": ["Pump", "Valve", "Pipeline", "Water Tank", "Drainage Line", "Fixture", "Pressure Vessel"],
        "readings": [
            ("Line pressure", "bar", "Site standard"), ("Flow rate", "LPM", "Design range"),
            ("Pump current", "A", "Nameplate +/-10%"), ("Tank level", "%", "Operational range"),
            ("Leak status", "Pass/Fail", "No leak"), ("Drainage flow", "Pass/Fail", "Clear"),
        ],
        "tasks": [
            "Inspect visible leaks and seepage", "Check isolation valves and fixtures",
            "Verify pump operation, noise and vibration", "Perform pressure/flow test where applicable",
            "Check tank cleanliness, overflow and float valve", "Clear blockages and flush line if required",
            "Record repair materials and post-repair observation",
        ],
        "risks": ["Water contamination", "Slip hazard", "Confined area", "Hot water/pressure release"],
    },
    "Electrical": {
        "assets": ["LT Panel", "DB", "MCC", "Breaker", "Switchgear", "UPS", "Inverter", "Generator", "Lighting Circuit"],
        "readings": [
            ("R-Y voltage", "V", "400-440"), ("Y-B voltage", "V", "400-440"), ("B-R voltage", "V", "400-440"),
            ("R phase current", "A", "Load design"), ("Y phase current", "A", "Load design"), ("B phase current", "A", "Load design"),
            ("Earth resistance", "ohm", "< 5 or site standard"), ("Insulation resistance", "Mohm", "> 1"),
            ("Neutral voltage", "V", "Low/within standard"),
        ],
        "tasks": [
            "Apply isolation and lockout/tagout", "Inspect panel cleanliness and ventilation",
            "Check breaker, contactor and terminal tightness", "Record voltage/current/earth/IR readings",
            "Check cable gland, labeling and warning stickers", "Verify emergency lighting/backup operation",
            "Restore supply and observe load stability",
        ],
        "risks": ["Electric shock", "Arc flash", "Stored energy", "Unauthorized energization"],
    },
    "Refrigeration": {
        "assets": ["Cold Room", "Compressor", "Evaporator", "Condenser", "Expansion Valve", "Refrigerant Line"],
        "readings": [
            ("Room temperature", "C", "Setpoint +/- tolerance"), ("Product temperature", "C", "Client spec"),
            ("Suction pressure", "psi/bar", "OEM range"), ("Discharge pressure", "psi/bar", "OEM range"),
            ("Superheat", "C", "OEM range"), ("Subcooling", "C", "OEM range"),
            ("Compressor current", "A", "Nameplate +/-10%"), ("Defrost cycle", "Pass/Fail", "Operational"),
        ],
        "tasks": [
            "Inspect compressor oil/noise/vibration", "Check evaporator and condenser condition",
            "Check refrigerant leak signs and insulation", "Record pressure/temperature readings",
            "Verify defrost operation and drain clearing", "Check door gasket and cold room panel condition",
            "Confirm final box temperature and alarm status",
        ],
        "risks": ["Refrigerant exposure", "Low temperature exposure", "Electrical isolation", "Food/product spoilage"],
    },
    "Fire Safety": {
        "assets": ["Extinguisher", "Hydrant", "Sprinkler", "Fire Alarm Panel", "Smoke Detector", "Pump Room", "Exit Sign"],
        "readings": [
            ("Extinguisher pressure", "bar/zone", "Green zone"), ("Hydrant pressure", "bar", "Code/site standard"),
            ("Jockey pump pressure", "bar", "Setpoint"), ("Main pump pressure", "bar", "Setpoint"),
            ("Detector response", "Pass/Fail", "Pass"), ("Alarm panel status", "Normal/Fault", "Normal"),
            ("Emergency light duration", "min", "Code/site standard"),
        ],
        "tasks": [
            "Verify access and physical condition", "Check pressure, seal, pin, tag and expiry",
            "Inspect hydrant/sprinkler valves and pressure", "Test alarm panel, detector and sounder status",
            "Check emergency lighting and exit signage", "Record defect and corrective action priority",
            "Capture compliance sign-off where required",
        ],
        "risks": ["Compliance failure", "Life safety impairment", "Blocked access", "Pump/electrical hazards"],
    },
}

GENERIC_PROFILE = {
    "assets": ["Asset", "Work Area", "Document", "Material", "Customer Account"],
    "readings": [("Status", "Pass/Fail", "Pass"), ("Quantity", "Nos", "As required"), ("Amount", "INR", "Approved")],
    "tasks": [
        "Verify request details", "Confirm approval and required documents", "Record observations",
        "Attach photo/document reference if applicable", "Capture responsible person sign-off",
    ],
    "risks": ["Approval gap", "Data mismatch", "Missing documentation"],
}

SAMPLE_FORMS = {
    "HVAC": "preventive maintenance checklist",
    "Electrical": "load/voltage reading sheet",
    "Fire Safety": "extinguisher inspection form",
    "Refrigeration": "cold room checklist",
    "Plumbing": "pressure test form",
    "Commercial Service Shared": "job card",
    "Inventory": "goods received note",
    "Finance Accounting": "invoice approval form",
}


def slug(text):
    out = []
    for ch in text.title().replace("/", " ").replace("&", "And"):
        if ch.isalnum():
            out.append(ch)
        elif ch.isspace() or ch in "-_":
            out.append("_")
    value = "".join(out)
    while "__" in value:
        value = value.replace("__", "_")
    return value.strip("_")


def profile(trade):
    return TRADE_PROFILES.get(trade, GENERIC_PROFILE)


def recommended_module(trade, form):
    f = form.lower()
    if "invoice" in f or "payment" in f or "credit" in f or "collections" in f or "costing" in f:
        return "Finance / Payments"
    if "purchase" in f or "goods received" in f or "stock" in f or "parts" in f or "inventory" in f:
        return "Inventory / Purchases"
    if "dispatch" in f or "attendance" in f or "leave" in f:
        return "Dispatch / Employees"
    if "quote" in f or "estimate" in f or "lead" in f or "enquiry" in f:
        return "Quotations / Clients"
    if "amc" in f or "warranty" in f:
        return "Contracts / Jobs"
    if "safety" in f or "permit" in f or "fire" in f or "ppe" in f or "risk" in f:
        return "Compliance / Service Desk"
    return "Jobs / Service Desk"


def flags(form, trade):
    f = form.lower()
    compliance_trade = trade in ("Fire Safety", "Compliance Safety", "Electrical")
    return {
        "customer_signature": "Yes" if any(x in f for x in ["sign-off", "approval", "handover", "completion", "service report", "receipt", "invoice", "job card"]) else "No",
        "technician_signature": "Yes" if any(x in f for x in ["checklist", "inspection", "service", "report", "test", "commissioning", "maintenance", "job card", "drill"]) else "No",
        "photo": "Yes" if any(x in f for x in ["inspection", "report", "defect", "damage", "site", "completion", "asset", "equipment", "fire", "leak"]) else "No",
        "readings": "Yes" if any(x in f for x in ["reading", "pressure", "temperature", "voltage", "load", "test", "earthing", "insulation", "refrigerant", "water treatment", "pump", "chiller", "cold room"]) else "No",
        "compliance": "Yes" if compliance_trade or any(x in f for x in ["fire", "safety", "noc", "audit", "lockout", "permit", "earthing", "insulation", "evacuation", "drill", "ppe", "risk"]) else "No",
    }


def fields_for(form, trade):
    base = [
        "Company Name", "Client", "Site", "Job Number", "Date", "Time", "Technician", "Equipment/Asset",
        "Asset Serial", "Job Type", "Priority", "Status", "Checklist Status", "Remarks",
    ]
    p = profile(trade)
    for reading, unit, _ in p["readings"][:8]:
        base.append(f"{reading} ({unit})")
    base += ["Parts Used", "Photo References", "Customer Signature", "Technician Signature"]
    return "; ".join(dict.fromkeys(base))


def workflow_trigger(trade, form):
    f = form.lower()
    if "amc" in f or "preventive" in f:
        return "AMC visit or scheduled maintenance job"
    if "breakdown" in f or "troubleshooting" in f or "complaint" in f:
        return "Service Desk incident converted to job"
    if "commissioning" in f or "installation" in f:
        return "New installation or project handover"
    if "parts" in f or "stock" in f or "goods received" in f:
        return "Inventory or purchase workflow"
    if "invoice" in f or "payment" in f or "credit" in f:
        return "Finance approval workflow"
    if trade in ("Fire Safety", "Compliance Safety"):
        return "Compliance inspection workflow"
    return "Field-service job workflow"


def checklist_items(trade, form):
    f = form.lower()
    p = profile(trade)
    items = list(p["tasks"])
    if "survey" in f or "site" in f:
        items += ["Record access constraints, power/water availability and installation feasibility", "Capture site photos and client requirements"]
    if "commissioning" in f or "installation" in f:
        items += ["Verify installation as per approved scope", "Run startup test and record acceptance readings"]
    if "cleaning" in f:
        items += ["Record before cleaning condition", "Confirm post-cleaning airflow/water flow and housekeeping"]
    if "quote" in f or "estimate" in f:
        items += ["Confirm scope, exclusions, material requirement and validity period"]
    if "sign-off" in f or "handover" in f:
        items += ["Confirm customer acceptance, pending points and next visit date"]
    return list(dict.fromkeys(items))


def validation_rules(trade, form):
    f = form.lower()
    rules = [
        {"field": "client", "required": True, "type": "text"},
        {"field": "site", "required": True, "type": "text"},
        {"field": "job_number", "required": True, "type": "text"},
        {"field": "technician", "required": True, "type": "text"},
        {"field": "date", "required": True, "type": "date"},
    ]
    fl = flags(form, trade)
    if fl["customer_signature"] == "Yes":
        rules.append({"field": "customer_signature", "required": True, "type": "signature", "complete_before": "close_job"})
    if fl["technician_signature"] == "Yes":
        rules.append({"field": "technician_signature", "required": True, "type": "signature"})
    if fl["photo"] == "Yes":
        rules.append({"field": "photo_references", "required": True, "type": "photo_list"})
    if fl["readings"] == "Yes":
        for reading, unit, normal in profile(trade)["readings"]:
            rules.append({"field": slug(reading).lower(), "required": False, "type": "reading", "unit": unit, "normal_range": normal})
    if "gst" in f or "invoice" in f or "credit" in f:
        rules.append({"field": "gstin", "required": False, "type": "gstin"})
    return rules


def schema_for(folder, trade, form, file_name):
    fl = flags(form, trade)
    return {
        "schemaVersion": VERSION,
        "templateId": slug(f"{trade}_{form}").lower(),
        "trade": trade,
        "categoryFolder": folder,
        "formName": form.title(),
        "defaultFileName": file_name,
        "recommendedServoERPModule": recommended_module(trade, form),
        "workflowTrigger": workflow_trigger(trade, form),
        "mobileSuitable": True,
        "signatureRequirements": {
            "customer": fl["customer_signature"] == "Yes",
            "technician": fl["technician_signature"] == "Yes",
        },
        "captureRequirements": {
            "photos": fl["photo"] == "Yes",
            "readings": fl["readings"] == "Yes",
            "complianceCritical": fl["compliance"] == "Yes",
        },
        "autofill": [
            "company_name", "client", "site", "job_number", "job_type", "priority", "technician", "scheduled_date", "asset_id", "asset_serial",
        ],
        "fields": build_schema_fields(trade, form),
        "validationRules": validation_rules(trade, form),
        "pdfExport": {
            "includeHeader": True,
            "includePhotos": True,
            "includeSignatures": True,
            "includeAuditTrail": True,
        },
        "versioning": {
            "templateVersion": VERSION,
            "retainCompletedFormVersion": True,
            "createdBy": "ServoERP Research Agent",
            "createdDate": DATE_ACCESSED,
        },
    }


def build_schema_fields(trade, form):
    fields = [
        {"name": "company_name", "label": "Company Name", "type": "text", "required": True},
        {"name": "client", "label": "Client", "type": "lookup", "source": "Clients", "required": True},
        {"name": "site", "label": "Site", "type": "lookup", "source": "ClientSites", "required": True},
        {"name": "job_number", "label": "Job Number", "type": "text", "required": True},
        {"name": "technician", "label": "Technician", "type": "lookup", "source": "Employees", "required": True},
        {"name": "equipment_asset", "label": "Equipment / Asset", "type": "lookup", "source": "ClientAssets", "required": False},
        {"name": "priority", "label": "Priority", "type": "select", "options": ["Low", "Normal", "Medium", "High", "Critical"], "required": True},
        {"name": "status", "label": "Status", "type": "select", "options": ["Open", "In Progress", "Completed", "Pending Customer", "Requires Follow-up"], "required": True},
    ]
    for item in checklist_items(trade, form):
        fields.append({"name": slug(item).lower(), "label": item, "type": "check_status", "options": ["OK", "Needs Action", "N/A"], "required": False})
    for reading, unit, normal in profile(trade)["readings"]:
        fields.append({"name": slug(reading).lower(), "label": reading, "type": "reading", "unit": unit, "normalRange": normal, "required": False})
    fields += [
        {"name": "parts_used", "label": "Parts Used", "type": "line_items", "columns": ["part_code", "description", "qty", "uom", "rate", "amount"], "required": False},
        {"name": "photo_references", "label": "Photo References", "type": "photos", "required": flags(form, trade)["photo"] == "Yes"},
        {"name": "technician_remarks", "label": "Technician Remarks", "type": "textarea", "required": False},
        {"name": "customer_signature", "label": "Customer Signature", "type": "signature", "required": flags(form, trade)["customer_signature"] == "Yes"},
        {"name": "technician_signature", "label": "Technician Signature", "type": "signature", "required": flags(form, trade)["technician_signature"] == "Yes"},
    ]
    return fields


def style_sheet(ws):
    colors = {
        "blue": "1D4ED8",
        "blue_dark": "0F172A",
        "blue_light": "EFF6FF",
        "green": "16A34A",
        "purple": "7C3AED",
        "amber": "F59E0B",
        "gray": "F8FAFC",
        "border": "D9E2EC",
        "red": "DC2626",
    }
    border = Border(
        left=Side(style="thin", color=colors["border"]),
        right=Side(style="thin", color=colors["border"]),
        top=Side(style="thin", color=colors["border"]),
        bottom=Side(style="thin", color=colors["border"]),
    )
    for row in ws.iter_rows():
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = Font(name="Segoe UI", size=10, color="0F172A")
    for merged in ["A1:L1", "A2:L2"]:
        pass
    ws["A1"].fill = PatternFill("solid", fgColor=colors["blue_dark"])
    ws["A1"].font = Font(name="Segoe UI", size=16, bold=True, color="FFFFFF")
    ws["A2"].fill = PatternFill("solid", fgColor=colors["blue_light"])
    ws["A2"].font = Font(name="Segoe UI", size=10, italic=True, color="334155")
    for row in range(4, ws.max_row + 1):
        first = str(ws.cell(row, 1).value or "")
        if first.startswith("SECTION:"):
            for col in range(1, 13):
                cell = ws.cell(row, col)
                cell.fill = PatternFill("solid", fgColor=colors["blue_light"])
                cell.font = Font(name="Segoe UI", size=11, bold=True, color=colors["blue"])
        elif row % 2 == 0:
            for col in range(1, min(13, ws.max_column + 1)):
                ws.cell(row, col).fill = PatternFill("solid", fgColor=colors["gray"])
    widths = [28, 20, 20, 16, 16, 16, 18, 18, 18, 24, 18, 24]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width


def create_xlsx(path, trade, form, sample=False):
    wb = Workbook()
    ws = wb.active
    ws.title = "Smart Form"
    p = profile(trade)
    title = f"ServoERP {trade} - {form.title()}"
    if sample:
        title += " (Completed Sample)"
    ws.merge_cells("A1:L1")
    ws["A1"] = title
    ws.merge_cells("A2:L2")
    ws["A2"] = f"Premium smart form template | Version {VERSION} | Workflow: {workflow_trigger(trade, form)}"
    ws.append([])
    ws.append(["SECTION: Job / Client / Site", "", "", "", "", "", "", "", "", "", "", ""])
    ws.append(["Company Name", "ServoERP", "Client", "ABC Cooling Solutions Pvt. Ltd." if sample else "", "Site", "Main Site - Pune" if sample else "", "Job Number", "JOB-2026-0001" if sample else "", "Date", DATE_ACCESSED if sample else "", "Technician", "Amit Sharma" if sample else ""])
    ws.append(["Equipment / Asset", p["assets"][0], "Asset Serial", "ASSET-001" if sample else "", "Priority", "Medium", "Status", "Completed" if sample else "Open", "Job Type", form.title(), "Photo Ref", "IMG-001, IMG-002" if sample else ""])
    ws.append([])
    ws.append(["SECTION: Safety / Compliance Gate", "", "", "", "", "", "", "", "", "", "", ""])
    for risk in p["risks"]:
        ws.append([risk, "Checked" if sample else "", "Risk Level", "Medium" if sample else "", "Control Applied", "Yes" if sample else "", "Permit Required", "No" if sample else "", "Remarks", "", "", ""])
    ws.append([])
    ws.append(["SECTION: Checklist", "", "", "", "", "", "", "", "", "", "", ""])
    ws.append(["Task", "OK", "Needs Action", "N/A", "Assigned To", "Due Date", "Photo Required", "Photo Ref", "Customer Visible", "Remarks", "Follow-Up", "Closed"])
    for idx, item in enumerate(checklist_items(trade, form), 1):
        ws.append([item, "Yes" if sample else "", "", "", "Technician" if sample else "", "", "Yes" if idx % 3 == 0 else "No", f"IMG-{idx:03d}" if sample and idx % 3 == 0 else "", "Yes", "Completed" if sample else "", "No", "Yes" if sample else ""])
    ws.append([])
    ws.append(["SECTION: Readings", "", "", "", "", "", "", "", "", "", "", ""])
    ws.append(["Reading", "Before", "After", "Unit", "Normal Range", "Result", "Instrument", "Instrument ID", "Photo Ref", "Remarks", "", ""])
    for idx, (reading, unit, normal) in enumerate(p["readings"], 1):
        before, after = sample_reading(reading, idx) if sample else ("", "")
        ws.append([reading, before, after, unit, normal, "OK" if sample else "", "Digital Meter" if sample else "", f"INS-{idx:03d}" if sample else "", "", "", "", ""])
    ws.append([])
    ws.append(["SECTION: Parts / Materials", "", "", "", "", "", "", "", "", "", "", ""])
    ws.append(["Part Code", "Description", "Qty", "UOM", "Rate", "Amount", "Issued From", "Returned Qty", "Warranty", "Remarks", "", ""])
    for row in sample_parts(trade) if sample else range(6):
        ws.append(row if sample else ["", "", "", "", "", "", "", "", "", "", "", ""])
    ws.append([])
    ws.append(["SECTION: Final Status / Sign-Off", "", "", "", "", "", "", "", "", "", "", ""])
    ws.append(["Completion Status", "Completed" if sample else "", "Customer Name", "Ravi Mehta" if sample else "", "Customer Signature", "[signature captured]" if sample else "", "Technician Signature", "[signature captured]" if sample else "", "Next Visit Due", "2026-06-10" if sample else "", "Follow-Up Required", "No" if sample else ""])
    ws.append(["Technician Remarks", sample_remark(trade, form) if sample else "", "", "", "", "", "", "", "", "", "", ""])
    style_sheet(ws)
    ws.freeze_panes = "A8"
    wb.save(path)


def sample_reading(reading, idx):
    text = reading.lower()
    if "temperature" in text:
        return ("29", "16")
    if "pressure" in text:
        return ("4.2", "4.4")
    if "voltage" in text:
        return ("415", "414")
    if "current" in text:
        return ("8.5", "8.2")
    if "resistance" in text:
        return ("2.1", "1.8")
    if "status" in text or "leak" in text or "defrost" in text:
        return ("Needs check", "Pass")
    return (str(idx), str(idx))


def sample_parts(trade):
    if trade == "HVAC":
        return [["FLT-24X24", "Return air filter", 2, "Nos", 650, 1300, "Main Store", 0, "No", "Replaced", "", ""]]
    if trade == "Electrical":
        return [["LUG-35", "Cable lug 35 sqmm", 6, "Nos", 45, 270, "Electrical Store", 0, "No", "Panel termination", "", ""]]
    if trade == "Fire Safety":
        return [["SEAL-FE", "Extinguisher safety seal", 4, "Nos", 15, 60, "Fire Store", 0, "No", "Seal replaced", "", ""]]
    return [["GEN-001", "Consumable item", 1, "Nos", 100, 100, "Store", 0, "No", "Used during service", "", ""]]


def sample_remark(trade, form):
    return f"{trade} {form} completed. Readings are within acceptable range. Customer briefed and sign-off captured."


def create_docx(path, trade, form, sample=False):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(9)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(f"ServoERP {trade} - {form.title()}" + (" (Completed Sample)" if sample else ""))
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(15, 23, 42)
    sub = doc.add_paragraph(f"Premium smart form | Version {VERSION} | Workflow: {workflow_trigger(trade, form)}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    info = doc.add_table(rows=4, cols=4)
    info.style = "Table Grid"
    rows = [
        ("Company", "ServoERP", "Job Number", "JOB-2026-0001" if sample else ""),
        ("Client", "ABC Cooling Solutions Pvt. Ltd." if sample else "", "Site", "Main Site - Pune" if sample else ""),
        ("Date", DATE_ACCESSED if sample else "", "Technician", "Amit Sharma" if sample else ""),
        ("Equipment / Asset", profile(trade)["assets"][0], "Priority / Status", "Medium / Completed" if sample else ""),
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            info.cell(r, c).text = value

    add_doc_table(doc, "Safety / Compliance Gate", ["Risk", "Checked", "Control", "Remarks"], [[risk, "Yes" if sample else "", "", ""] for risk in profile(trade)["risks"]])
    add_doc_table(doc, "Checklist", ["Task", "OK", "Needs Action", "N/A", "Remarks"], [[item, "Yes" if sample else "", "", "", "Completed" if sample else ""] for item in checklist_items(trade, form)])
    add_doc_table(doc, "Readings", ["Reading", "Before", "After", "Unit", "Normal Range"], [[r, *(sample_reading(r, i) if sample else ("", "")), u, n] for i, (r, u, n) in enumerate(profile(trade)["readings"], 1)])
    add_doc_table(doc, "Parts Used", ["Part Code", "Description", "Qty", "UOM", "Remarks"], [[row[0], row[1], str(row[2]), row[3], row[9]] for row in (sample_parts(trade) if sample else [["", "", "", "", "", "", "", "", "", "", "", ""] for _ in range(4)])])

    doc.add_heading("Remarks and Sign-Off", level=2)
    doc.add_paragraph(sample_remark(trade, form) if sample else "Technician remarks:")
    sign = doc.add_table(rows=2, cols=4)
    sign.style = "Table Grid"
    for i, header in enumerate(["Customer Name", "Customer Signature", "Technician Signature", "Date/Time"]):
        sign.cell(0, i).text = header
        sign.cell(1, i).text = ["Ravi Mehta", "[signature captured]", "[signature captured]", DATE_ACCESSED][i] if sample else ""
    doc.save(path)


def add_doc_table(doc, title, headers, rows):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values[:len(headers)]):
            cells[i].text = str(value)


def write_json(path, data):
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def build():
    if ROOT.exists():
        shutil.rmtree(ROOT)
    ROOT.mkdir(parents=True)
    schema_root = ROOT / "11_ServoERP_Original_Templates" / "Smart_Form_Schemas"
    sample_root = ROOT / "11_ServoERP_Original_Templates" / "Completed_Samples"
    schema_root.mkdir(parents=True)
    sample_root.mkdir(parents=True)

    mapping_rows = []
    created_rows = []
    validation_rows = []
    total_templates = 0
    total_schemas = 0
    total_samples = 0

    for folder, trade, forms in CATEGORIES:
        category_dir = ROOT / folder
        category_dir.mkdir(parents=True, exist_ok=True)
        for form in forms:
            base = f"{slug(trade)}_{slug(form)}_Premium"
            xlsx_name = base + ".xlsx"
            xlsx_path = category_dir / xlsx_name
            create_xlsx(xlsx_path, trade, form)
            total_templates += 1
            fl = flags(form, trade)
            mapping_rows.append([
                folder, trade, form.title(), xlsx_name, "Original Premium", "", recommended_module(trade, form),
                fields_for(form, trade), "Yes", fl["customer_signature"], fl["technician_signature"],
                fl["photo"], fl["readings"], fl["compliance"], VERSION, workflow_trigger(trade, form),
            ])
            created_rows.append([trade, form.title(), xlsx_name, "Premium original generated with trade-specific sections, readings, validation metadata, and sign-off areas.", f"{trade} {form} template PDF XLSX public"])

            schema = schema_for(folder, trade, form, xlsx_name)
            schema_name = base + ".schema.json"
            write_json(schema_root / schema_name, schema)
            total_schemas += 1
            for rule in schema["validationRules"]:
                validation_rows.append([schema["templateId"], trade, form.title(), rule.get("field"), rule.get("type"), rule.get("required"), rule.get("normal_range", ""), rule.get("complete_before", "")])

            if any(key in form.lower() for key in ["sign-off", "report", "certificate", "quotation", "estimate", "invoice", "work order", "job card", "approval", "checklist"]):
                docx_name = base + ".docx"
                create_docx(category_dir / docx_name, trade, form)
                total_templates += 1
                mapping_rows.append([
                    folder, trade, form.title(), docx_name, "Original Premium", "", recommended_module(trade, form),
                    fields_for(form, trade), "Yes", fl["customer_signature"], fl["technician_signature"],
                    fl["photo"], fl["readings"], fl["compliance"], VERSION, workflow_trigger(trade, form),
                ])
                created_rows.append([trade, form.title(), docx_name, "Premium DOCX companion generated for print/sign-off workflow.", f"{trade} {form} docx template public"])

            if SAMPLE_FORMS.get(trade) == form:
                sample_name = base + "_Completed_Sample.xlsx"
                create_xlsx(sample_root / sample_name, trade, form, sample=True)
                total_samples += 1
                sample_docx = base + "_Completed_Sample.docx"
                create_docx(sample_root / sample_docx, trade, form, sample=True)
                total_samples += 1

    write_mapping(mapping_rows)
    write_validation_index(validation_rows)
    write_sources()
    write_download_log()
    write_created(created_rows)
    write_readme(total_templates, total_schemas, total_samples, len(SEARCH_REFERENCES))
    write_manifest(mapping_rows, total_templates, total_schemas, total_samples)

    bad = validate_files()
    if bad:
        raise RuntimeError("Empty files created: " + "; ".join(bad))
    create_zip()
    return total_templates, total_schemas, total_samples, len(mapping_rows)


def write_mapping(rows):
    with (ROOT / "ServoERP_Form_Template_Mapping.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Category", "Trade", "FormName", "FileName", "SourceType", "SourceURL",
            "RecommendedServoERPModule", "RecommendedFields", "MobileSuitableYesNo",
            "RequiresCustomerSignatureYesNo", "RequiresTechnicianSignatureYesNo",
            "RequiresPhotoUploadYesNo", "RequiresReadingsYesNo", "ComplianceCriticalYesNo",
            "TemplateVersion", "WorkflowTrigger",
        ])
        writer.writerows(rows)


def write_validation_index(rows):
    with (ROOT / "ServoERP_Form_Template_Validation_Rules.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["TemplateId", "Trade", "FormName", "Field", "Type", "Required", "NormalRange", "CompletionGate"])
        writer.writerows(rows)


def write_sources():
    with (ROOT / "SOURCES.md").open("w", encoding="utf-8") as f:
        f.write("# Sources\n\n")
        f.write("No third-party template file was copied into this package. Links below were used only for research awareness and legal/source evaluation.\n\n")
        f.write("| File Name | Source URL | License / Usage Notes | Date Accessed |\n")
        f.write("|---|---|---|---|\n")
        for query, url, note in SEARCH_REFERENCES:
            f.write(f"| Research only: {query} | {url or 'N/A'} | {note} | {DATE_ACCESSED} |\n")


def write_download_log():
    with (ROOT / "DOWNLOAD_LOG.md").open("w", encoding="utf-8") as f:
        f.write("# Download Log\n\n")
        f.write("| Attempted Query | Downloaded Files | Skipped Links | Reason Skipped |\n")
        f.write("|---|---|---|---|\n")
        for query, url, note in SEARCH_REFERENCES:
            f.write(f"| {query} | None | {url or 'Search results only'} | {note} |\n")


def write_created(rows):
    with (ROOT / "MISSING_OR_CREATED_FORMS.md").open("w", encoding="utf-8") as f:
        f.write("# Missing Or Created Forms\n\n")
        f.write("All templates are original premium ServoERP-compatible files. Public references were not copied because consistent clean legal/unbranded direct downloads were not available for every requested form.\n\n")
        f.write("| Trade | Form Name | File Name | Reason | Recommended Future Search Query |\n")
        f.write("|---|---|---|---|---|\n")
        for row in rows:
            f.write("| " + " | ".join(str(x).replace("|", "/") for x in row) + " |\n")


def write_readme(total_templates, total_schemas, total_samples, skipped):
    summary = {}
    for folder, trade, _ in CATEGORIES:
        summary[trade] = len(list((ROOT / folder).glob("*.*")))
    with (ROOT / "README_INDEX.md").open("w", encoding="utf-8") as f:
        f.write("# ServoERP Premium Field-Service Smart Form Library\n\n")
        f.write("Premium original editable forms for HVAC, Plumbing, Electrical, Refrigeration, Fire Safety, Commercial Service, Dispatch, Inventory, Accounting, and Compliance workflows.\n\n")
        f.write("## What Makes This Premium\n\n")
        f.write("- Trade-specific readings, normal ranges, safety gates, parts tables, photo references, and signature blocks.\n")
        f.write("- JSON smart-form schemas for mobile/dynamic form rendering.\n")
        f.write("- Validation metadata for required fields, signatures, readings, GSTIN-style fields, photo capture, and completion gates.\n")
        f.write("- Filled sample forms for high-value workflows.\n")
        f.write("- Template versioning metadata for future completed-form retention.\n\n")
        f.write("## Folder Structure\n\n")
        for folder, trade, _ in CATEGORIES:
            f.write(f"- `{folder}` - {trade}\n")
        f.write("- `11_ServoERP_Original_Templates/Smart_Form_Schemas` - JSON schemas\n")
        f.write("- `11_ServoERP_Original_Templates/Completed_Samples` - filled examples\n\n")
        f.write("## Counts\n\n")
        f.write(f"- Premium editable templates: {total_templates}\n")
        f.write(f"- Smart-form JSON schemas: {total_schemas}\n")
        f.write(f"- Completed sample files: {total_samples}\n")
        f.write("- Downloaded public files: 0\n")
        f.write(f"- Skipped third-party files: {skipped}\n")
        f.write(f"- Template version: {VERSION}\n\n")
        f.write("## Category Summary\n\n")
        for trade, count in summary.items():
            f.write(f"- {trade}: {count} editable template files\n")
        f.write("\n## How ServoERP Can Use These Forms\n\n")
        f.write("- Jobs / Service Desk: attach suggested forms by job type or incident category.\n")
        f.write("- Contracts: generate AMC visit checklists and preventive maintenance reports.\n")
        f.write("- Inventory / Purchases: use requisition, stock issue, GRN, and parts-used forms.\n")
        f.write("- Finance: approval, receipt, credit note, collections, and job-costing workflows.\n")
        f.write("- Mobile: render schemas into app forms with photos, readings, and signatures.\n")


def write_manifest(mapping_rows, total_templates, total_schemas, total_samples):
    manifest = {
        "library": "ServoERP Premium Field-Service Smart Form Library",
        "version": VERSION,
        "createdDate": DATE_ACCESSED,
        "editableTemplateCount": total_templates,
        "schemaCount": total_schemas,
        "completedSampleCount": total_samples,
        "sourcePolicy": "Original ServoERP-compatible forms only; no copied branded/proprietary templates.",
        "recommendedNextStep": "Import JSON schemas into ServoERP dynamic form renderer and attach templates to Jobs/Service Desk/Contracts.",
        "trades": sorted(set(row[1] for row in mapping_rows)),
    }
    write_json(ROOT / "ServoERP_Form_Library_Manifest.json", manifest)


def validate_files():
    bad = []
    for path in ROOT.rglob("*"):
        if path.is_file() and path.stat().st_size <= 0:
            bad.append(str(path))
    # Open-read a few important formats.
    for path in list(ROOT.rglob("*.xlsx"))[:3]:
        load_workbook(path, read_only=True).close()
    for path in list(ROOT.rglob("*.schema.json"))[:3]:
        json.loads(path.read_text(encoding="utf-8"))
    return bad


def create_zip():
    ZIP_PATH.parent.mkdir(parents=True, exist_ok=True)
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in ROOT.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(ROOT.parent))


if __name__ == "__main__":
    templates, schemas, samples, mapping = build()
    print(f"Premium editable templates: {templates}")
    print(f"Smart-form schemas: {schemas}")
    print(f"Completed sample files: {samples}")
    print(f"Mapping rows: {mapping}")
    print(f"Root: {ROOT}")
    print(f"Zip: {ZIP_PATH}")
